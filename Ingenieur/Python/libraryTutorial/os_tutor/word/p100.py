from docx import Document

doc = Document()

# 添加标题
doc.add_heading('标题0', 0)
doc.add_heading('标题1', 1)
doc.add_heading('标题2', 2)
doc.add_heading('标题3', 3)
doc.save('/Users/myry/Documents/MyPythonProject/lernLibs/os_tutor/word/head01.docx')