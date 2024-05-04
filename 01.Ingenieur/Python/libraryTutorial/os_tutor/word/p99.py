from docx import Document

# document = Document() # 创建一个空的docx文件
# document.save('/Users/myry/Documents/MyPythonProject/lernLibs/os_tutor/word/doc01.docx') # 保存

document = Document('/Users/myry/Documents/MyPythonProject/lernLibs/os_tutor/word/doc01.docx')

# 添加段落
paragraph1 = document.add_paragraph('我要学习办公自动化')
paragraph2 = document.add_paragraph('正在学习word操作')

# 插入段落
paragraph3 = paragraph1.insert_paragraph_before('今天开始学习')
document.save('/Users/myry/Documents/MyPythonProject/lernLibs/os_tutor/word/doc02.docx')